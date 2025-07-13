"""
Report generation system for DocBridgeGuard 2.0
"""

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Union

from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import pandas as pd

from ..models import ProcessingResult, ComparisonResult, ComplianceMetadata

logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    Generates various types of reports for DocBridgeGuard processing results
    """
    
    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize report generator
        
        Args:
            output_dir: Directory for generated reports
        """
        self.output_dir = output_dir or Path("reports")
        self.output_dir.mkdir(exist_ok=True)
        logger.info(f"Report generator initialized with output dir: {self.output_dir}")
    
    def generate_processing_report(
        self,
        result: ProcessingResult,
        format_type: str = "json"
    ) -> Path:
        """
        Generate report for a single processing result
        
        Args:
            result: Processing result to report on
            format_type: Report format ('json', 'pdf', 'docx')
            
        Returns:
            Path to generated report file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"processing_report_{result.document_id}_{timestamp}"
        
        if format_type == "json":
            return self._generate_json_report(result, filename)
        elif format_type == "pdf":
            return self._generate_pdf_processing_report(result, filename)
        elif format_type == "docx":
            return self._generate_docx_processing_report(result, filename)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def generate_comparison_report(
        self,
        comparison: ComparisonResult,
        format_type: str = "json"
    ) -> Path:
        """
        Generate report for provider comparison
        
        Args:
            comparison: Comparison result to report on
            format_type: Report format ('json', 'pdf', 'docx')
            
        Returns:
            Path to generated report file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"comparison_report_{comparison.document_id}_{timestamp}"
        
        if format_type == "json":
            return self._generate_json_comparison_report(comparison, filename)
        elif format_type == "pdf":
            return self._generate_pdf_comparison_report(comparison, filename)
        elif format_type == "docx":
            return self._generate_docx_comparison_report(comparison, filename)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def generate_compliance_report(
        self,
        results: List[ProcessingResult],
        format_type: str = "pdf"
    ) -> Path:
        """
        Generate compliance overview report for multiple results
        
        Args:
            results: List of processing results
            format_type: Report format ('json', 'pdf', 'docx')
            
        Returns:
            Path to generated report file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"compliance_report_{timestamp}"
        
        if format_type == "json":
            return self._generate_json_compliance_report(results, filename)
        elif format_type == "pdf":
            return self._generate_pdf_compliance_report(results, filename)
        elif format_type == "docx":
            return self._generate_docx_compliance_report(results, filename)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _generate_json_report(self, result: ProcessingResult, filename: str) -> Path:
        """Generate JSON processing report"""
        output_path = self.output_dir / f"{filename}.json"
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result.model_dump(), f, indent=2, default=str)
            
            logger.info(f"JSON report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate JSON report: {e}")
            raise
    
    def _generate_json_comparison_report(self, comparison: ComparisonResult, filename: str) -> Path:
        """Generate JSON comparison report"""
        output_path = self.output_dir / f"{filename}.json"
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(comparison.model_dump(), f, indent=2, default=str)
            
            logger.info(f"JSON comparison report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate JSON comparison report: {e}")
            raise
    
    def _generate_json_compliance_report(self, results: List[ProcessingResult], filename: str) -> Path:
        """Generate JSON compliance report"""
        output_path = self.output_dir / f"{filename}.json"
        
        try:
            # Aggregate compliance data
            compliance_data = {
                "report_timestamp": datetime.now().isoformat(),
                "total_documents": len(results),
                "compliance_summary": self._analyze_compliance_metrics(results),
                "document_details": [
                    {
                        "document_id": r.document_id,
                        "filename": r.original_filename,
                        "compliance_score": r.compliance_metadata.compliance_score,
                        "redactions_count": r.compliance_metadata.redactions_count,
                        "risk_flags": r.compliance_metadata.risk_flags,
                        "jurisdiction": r.compliance_metadata.jurisdiction.value,
                        "document_type": r.compliance_metadata.document_type.value
                    }
                    for r in results
                ]
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(compliance_data, f, indent=2, default=str)
            
            logger.info(f"JSON compliance report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate JSON compliance report: {e}")
            raise
    
    def _generate_pdf_processing_report(self, result: ProcessingResult, filename: str) -> Path:
        """Generate PDF processing report"""
        output_path = self.output_dir / f"{filename}.pdf"
        
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=16)
            
            # Title
            pdf.cell(200, 10, txt="DocBridgeGuard 2.0 - Processing Report", ln=True, align='C')
            pdf.ln(10)
            
            # Document information
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt=f"Document ID: {result.document_id}", ln=True)
            pdf.cell(200, 10, txt=f"Filename: {result.original_filename}", ln=True)
            pdf.cell(200, 10, txt=f"Provider: {result.provider_used}", ln=True)
            pdf.cell(200, 10, txt=f"Status: {result.status.value}", ln=True)
            pdf.cell(200, 10, txt=f"Processing Time: {result.processing_time_seconds:.2f}s", ln=True)
            pdf.ln(5)
            
            # Compliance information
            pdf.set_font("Arial", "B", 14)
            pdf.cell(200, 10, txt="Compliance Information", ln=True)
            pdf.set_font("Arial", size=12)
            
            metadata = result.compliance_metadata
            pdf.cell(200, 10, txt=f"Document Type: {metadata.document_type.value}", ln=True)
            pdf.cell(200, 10, txt=f"Jurisdiction: {metadata.jurisdiction.value}", ln=True)
            pdf.cell(200, 10, txt=f"Compliance Score: {metadata.compliance_score:.2f}", ln=True)
            pdf.cell(200, 10, txt=f"Redactions Applied: {metadata.redactions_count}", ln=True)
            pdf.cell(200, 10, txt=f"Legal Basis: {metadata.legal_basis}", ln=True)
            pdf.ln(5)
            
            # Bridge extraction summary
            pdf.set_font("Arial", "B", 14)
            pdf.cell(200, 10, txt="Bridge Extraction Summary", ln=True)
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt=f"Total Bridges: {len(result.bridges)}", ln=True)
            
            if result.bridges:
                privacy_counts = {}
                for bridge in result.bridges:
                    impact = bridge.privacy_impact.value
                    privacy_counts[impact] = privacy_counts.get(impact, 0) + 1
                
                for impact, count in privacy_counts.items():
                    pdf.cell(200, 10, txt=f"  {impact.title()} Impact: {count}", ln=True)
            
            pdf.ln(5)
            
            # Risk flags
            if metadata.risk_flags:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(200, 10, txt="Risk Flags", ln=True)
                pdf.set_font("Arial", size=12)
                for flag in metadata.risk_flags:
                    pdf.cell(200, 10, txt=f"  • {flag}", ln=True)
            
            pdf.output(str(output_path))
            logger.info(f"PDF report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate PDF report: {e}")
            raise
    
    def _generate_pdf_comparison_report(self, comparison: ComparisonResult, filename: str) -> Path:
        """Generate PDF comparison report"""
        output_path = self.output_dir / f"{filename}.pdf"
        
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=16)
            
            # Title
            pdf.cell(200, 10, txt="DocBridgeGuard 2.0 - Provider Comparison Report", ln=True, align='C')
            pdf.ln(10)
            
            # Comparison summary
            pdf.set_font("Arial", "B", 14)
            pdf.cell(200, 10, txt="Comparison Summary", ln=True)
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt=f"Document ID: {comparison.document_id}", ln=True)
            pdf.cell(200, 10, txt=f"Winner: {comparison.winner or 'Tie'}", ln=True)
            pdf.cell(200, 10, txt=f"Confidence: {comparison.confidence_in_winner:.2f}", ln=True)
            pdf.ln(5)
            
            # Metrics comparison
            pdf.set_font("Arial", "B", 14)
            pdf.cell(200, 10, txt="Key Metrics", ln=True)
            pdf.set_font("Arial", size=12)
            
            metrics = comparison.comparison_metrics
            for key, value in metrics.items():
                if isinstance(value, (int, float)):
                    pdf.cell(200, 10, txt=f"{key}: {value:.3f}", ln=True)
            
            pdf.ln(5)
            
            # Recommendations
            analysis = comparison.detailed_analysis
            if "recommendations" in analysis:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(200, 10, txt="Recommendations", ln=True)
                pdf.set_font("Arial", size=12)
                
                for recommendation in analysis["recommendations"]:
                    # Handle long text by splitting into multiple lines
                    words = recommendation.split()
                    line = ""
                    for word in words:
                        if len(line + word) < 80:  # Approximate character limit per line
                            line += word + " "
                        else:
                            pdf.cell(200, 10, txt=f"  • {line.strip()}", ln=True)
                            line = word + " "
                    if line:
                        pdf.cell(200, 10, txt=f"  • {line.strip()}", ln=True)
            
            pdf.output(str(output_path))
            logger.info(f"PDF comparison report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate PDF comparison report: {e}")
            raise
    
    def _generate_docx_processing_report(self, result: ProcessingResult, filename: str) -> Path:
        """Generate Word document processing report"""
        output_path = self.output_dir / f"{filename}.docx"
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading("DocBridgeGuard 2.0 - Processing Report", 0)
            
            # Document information
            doc.add_heading("Document Information", level=1)
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ("Document ID", result.document_id),
                ("Filename", result.original_filename),
                ("Provider", result.provider_used),
                ("Status", result.status.value),
                ("Processing Time", f"{result.processing_time_seconds:.2f} seconds")
            ]
            
            for i, (key, value) in enumerate(info_data):
                info_table.cell(i, 0).text = key
                info_table.cell(i, 1).text = str(value)
            
            # Compliance section
            doc.add_heading("Compliance Information", level=1)
            metadata = result.compliance_metadata
            
            compliance_table = doc.add_table(rows=5, cols=2)
            compliance_table.style = 'Table Grid'
            
            compliance_data = [
                ("Document Type", metadata.document_type.value),
                ("Jurisdiction", metadata.jurisdiction.value),
                ("Compliance Score", f"{metadata.compliance_score:.2f}"),
                ("Redactions Applied", str(metadata.redactions_count)),
                ("Legal Basis", metadata.legal_basis)
            ]
            
            for i, (key, value) in enumerate(compliance_data):
                compliance_table.cell(i, 0).text = key
                compliance_table.cell(i, 1).text = value
            
            # Bridge extraction
            doc.add_heading("Bridge Extraction Results", level=1)
            doc.add_paragraph(f"Total bridges extracted: {len(result.bridges)}")
            
            if result.bridges:
                # Privacy impact distribution
                privacy_counts = {}
                for bridge in result.bridges:
                    impact = bridge.privacy_impact.value
                    privacy_counts[impact] = privacy_counts.get(impact, 0) + 1
                
                doc.add_paragraph("Privacy Impact Distribution:")
                for impact, count in privacy_counts.items():
                    doc.add_paragraph(f"  • {impact.title()}: {count}", style='List Bullet')
            
            # Risk flags
            if metadata.risk_flags:
                doc.add_heading("Risk Flags", level=1)
                for flag in metadata.risk_flags:
                    doc.add_paragraph(f"• {flag}", style='List Bullet')
            
            doc.save(str(output_path))
            logger.info(f"DOCX report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX report: {e}")
            raise
    
    def _generate_docx_comparison_report(self, comparison: ComparisonResult, filename: str) -> Path:
        """Generate Word document comparison report"""
        output_path = self.output_dir / f"{filename}.docx"
        
        try:
            doc = Document()
            
            # Title
            doc.add_heading("DocBridgeGuard 2.0 - Provider Comparison Report", 0)
            
            # Summary
            doc.add_heading("Comparison Summary", level=1)
            summary_table = doc.add_table(rows=3, cols=2)
            summary_table.style = 'Table Grid'
            
            summary_data = [
                ("Document ID", comparison.document_id),
                ("Winner", comparison.winner or "Tie"),
                ("Confidence", f"{comparison.confidence_in_winner:.2f}")
            ]
            
            for i, (key, value) in enumerate(summary_data):
                summary_table.cell(i, 0).text = key
                summary_table.cell(i, 1).text = value
            
            # Detailed metrics
            doc.add_heading("Detailed Metrics", level=1)
            metrics = comparison.comparison_metrics
            
            for key, value in metrics.items():
                if isinstance(value, (int, float)):
                    doc.add_paragraph(f"{key}: {value:.3f}")
            
            # Recommendations
            analysis = comparison.detailed_analysis
            if "recommendations" in analysis:
                doc.add_heading("Recommendations", level=1)
                for recommendation in analysis["recommendations"]:
                    doc.add_paragraph(f"• {recommendation}", style='List Bullet')
            
            doc.save(str(output_path))
            logger.info(f"DOCX comparison report generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX comparison report: {e}")
            raise
    
    def _analyze_compliance_metrics(self, results: List[ProcessingResult]) -> Dict[str, Any]:
        """Analyze compliance metrics across multiple results"""
        if not results:
            return {}
        
        scores = [r.compliance_metadata.compliance_score for r in results]
        redactions = [r.compliance_metadata.redactions_count for r in results]
        
        # Document type distribution
        doc_types = {}
        jurisdictions = {}
        
        for result in results:
            doc_type = result.compliance_metadata.document_type.value
            jurisdiction = result.compliance_metadata.jurisdiction.value
            
            doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
            jurisdictions[jurisdiction] = jurisdictions.get(jurisdiction, 0) + 1
        
        # Risk analysis
        all_risk_flags = []
        for result in results:
            all_risk_flags.extend(result.compliance_metadata.risk_flags)
        
        risk_flag_counts = {}
        for flag in all_risk_flags:
            risk_flag_counts[flag] = risk_flag_counts.get(flag, 0) + 1
        
        return {
            "average_compliance_score": sum(scores) / len(scores),
            "min_compliance_score": min(scores),
            "max_compliance_score": max(scores),
            "total_redactions": sum(redactions),
            "average_redactions": sum(redactions) / len(redactions),
            "document_type_distribution": doc_types,
            "jurisdiction_distribution": jurisdictions,
            "top_risk_flags": dict(sorted(risk_flag_counts.items(), key=lambda x: x[1], reverse=True)[:10])
        }
    
    def export_to_excel(self, results: List[ProcessingResult], filename: str) -> Path:
        """Export results to Excel spreadsheet"""
        output_path = self.output_dir / f"{filename}.xlsx"
        
        try:
            # Create DataFrame from results
            data = []
            for result in results:
                row = {
                    "document_id": result.document_id,
                    "filename": result.original_filename,
                    "provider": result.provider_used,
                    "status": result.status.value,
                    "processing_time": result.processing_time_seconds,
                    "extracted_text_length": len(result.extracted_text),
                    "bridge_count": len(result.bridges),
                    "table_count": len(result.tables),
                    "document_type": result.compliance_metadata.document_type.value,
                    "jurisdiction": result.compliance_metadata.jurisdiction.value,
                    "compliance_score": result.compliance_metadata.compliance_score,
                    "redactions_count": result.compliance_metadata.redactions_count,
                    "risk_flags_count": len(result.compliance_metadata.risk_flags)
                }
                data.append(row)
            
            df = pd.DataFrame(data)
            df.to_excel(output_path, index=False)
            
            logger.info(f"Excel export generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate Excel export: {e}")
            raise